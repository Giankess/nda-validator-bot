from fastapi import UploadFile
from docx import Document
import uuid
import os
from typing import Dict, Any

class DocumentService:
    def __init__(self):
        self.documents_dir = "documents"
        os.makedirs(self.documents_dir, exist_ok=True)

    async def parse_document(self, file: UploadFile) -> str:
        """Parse the uploaded document and save it to disk."""
        document_id = str(uuid.uuid4())
        file_path = os.path.join(self.documents_dir, f"{document_id}.docx")
        
        # Save the uploaded file
        with open(file_path, "wb") as buffer:
            content = await file.read()
            buffer.write(content)
        
        return document_id

    async def get_document(self, document_id: str) -> Document:
        """Retrieve a document by its ID."""
        file_path = os.path.join(self.documents_dir, f"{document_id}.docx")
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Document {document_id} not found")
        return Document(file_path)

    async def create_redline_document(self, document: Document, suggestions: Dict[str, Any]) -> str:
        """Create a redline version of the document with suggested changes."""
        redline_id = str(uuid.uuid4())
        redline_doc = Document()
        
        for paragraph in document.paragraphs:
            if paragraph.text in suggestions:
                # Add the original text with strikethrough
                p = redline_doc.add_paragraph()
                run = p.add_run(paragraph.text)
                run.font.strike = True
                
                # Add the suggested text
                p = redline_doc.add_paragraph()
                run = p.add_run(suggestions[paragraph.text]["suggestion"])
                run.font.color.rgb = (0, 128, 0)  # Green color
            else:
                # Copy the original paragraph
                p = redline_doc.add_paragraph(paragraph.text)
        
        # Save the redline document
        redline_path = os.path.join(self.documents_dir, f"{redline_id}_redline.docx")
        redline_doc.save(redline_path)
        return redline_id

    async def create_clean_document(self, document_id: str) -> str:
        """Create a clean version of the document with accepted changes."""
        clean_id = str(uuid.uuid4())
        doc = await self.get_document(document_id)
        clean_doc = Document()
        
        for paragraph in doc.paragraphs:
            if paragraph.text.startswith("Suggested:"):
                # Skip the original text and only keep the suggestion
                continue
            clean_doc.add_paragraph(paragraph.text)
        
        # Save the clean document
        clean_path = os.path.join(self.documents_dir, f"{clean_id}_clean.docx")
        clean_doc.save(clean_path)
        return clean_id

    async def return_document(self, document_id: str) -> bytes:
        """Return the document as bytes for download."""
        file_path = os.path.join(self.documents_dir, f"{document_id}.docx")
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Document {document_id} not found")
        
        with open(file_path, "rb") as file:
            return file.read() 